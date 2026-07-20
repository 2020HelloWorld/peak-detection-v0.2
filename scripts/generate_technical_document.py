"""Generate the standalone algorithm implementation/provenance Word document."""

from __future__ import annotations

from pathlib import Path
import sys


PROJECT_ROOT = Path(__file__).resolve().parents[1]
VENDOR = PROJECT_ROOT / ".vendor"
if VENDOR.exists():
    sys.path.insert(0, str(VENDOR))

from docx import Document

from generate_deliverables import markdown_to_docx


SOURCE = PROJECT_ROOT / "docs" / "算法实现与技术溯源_v0.5.2.md"
DESTINATION = PROJECT_ROOT / "docs" / "算法实现与技术溯源_v0.5.2.docx"


def main() -> int:
    markdown_to_docx(SOURCE, DESTINATION)
    document = Document(DESTINATION)
    document.core_properties.title = "ChromPeak 算法实现与技术溯源说明"
    document.core_properties.subject = "色谱峰识别、噪声预处理、开源算法来源及代码定位"
    document.core_properties.author = "ChromPeak 项目"
    header = document.sections[0].header.paragraphs[0]
    header.text = "ChromPeak 算法实现与技术溯源 · v0.5.2"
    document.save(DESTINATION)
    print(
        {
            "source": str(SOURCE),
            "destination": str(DESTINATION),
            "bytes": DESTINATION.stat().st_size,
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
        }
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
